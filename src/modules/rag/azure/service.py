"""Azure AI Search RAG implementation for dooers_file_search_tool."""

from __future__ import annotations

import csv
import io
import json
import logging
import re
import uuid
from typing import Any

from azure.core.exceptions import ResourceNotFoundError

from src.config import settings as app_settings
from src.modules.external.azure.ai_search import (
    delete_documents,
    ensure_agent_index,
    get_document,
    search_documents,
    upload_documents,
)
from src.modules.rag.archive_storage import delete_archive_uri, upload_archive_bytes
from src.modules.rag.knowledge_repository import knowledge_repository
from src.modules.rag.knowledge_settings import normalize_knowledge_field_id

logger = logging.getLogger(__name__)


def _safe_index_name(agent_id: str, field_id: str) -> str:
    cleaned = re.sub(r"[^a-z0-9-]", "-", f"{agent_id}-{field_id}".lower())
    cleaned = re.sub(r"-+", "-", cleaned).strip("-")
    return cleaned[:64] or "agent"


# Searchable Edm.String fields are limited (~32 KiB UTF-8 per value); larger bodies get 207 + failed rows.
_AZURE_SEARCHABLE_CONTENT_MAX_BYTES = 32700


def _truncate_content_for_azure_index(content: str, *, filename: str) -> str:
    encoded = content.encode("utf-8")
    if len(encoded) <= _AZURE_SEARCHABLE_CONTENT_MAX_BYTES:
        return content
    suffix = "\n\n[Truncated: Azure AI Search limits searchable text to ~32 KiB per field.]"
    suf_b = suffix.encode("utf-8")
    budget = _AZURE_SEARCHABLE_CONTENT_MAX_BYTES - len(suf_b)
    if budget <= 0:
        return encoded[:_AZURE_SEARCHABLE_CONTENT_MAX_BYTES].decode("utf-8", errors="ignore")
    prefix = encoded[:budget].decode("utf-8", errors="ignore")
    combined = prefix + suffix
    while len(combined.encode("utf-8")) > _AZURE_SEARCHABLE_CONTENT_MAX_BYTES and prefix:
        prefix = prefix[:-1]
        combined = prefix + suffix
    logger.info(
        "Azure RAG: truncated extracted text for %r from %s to %s bytes (service field limit)",
        filename,
        len(encoded),
        len(combined.encode("utf-8")),
    )
    return combined


def _bytes_to_text(data: bytes, mime_type: str | None, filename: str) -> str:
    ext = "." + filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    def _decode_text(raw: bytes) -> str:
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return raw.decode(enc).strip()
            except Exception:
                continue
        return raw.decode("utf-8", errors="ignore").strip()

    if "text/" in (mime_type or "") or ext in {".txt", ".md"}:
        text = _decode_text(data)
        return text or f"[empty text document: {filename}]"

    if ext == ".json":
        try:
            obj = json.loads(_decode_text(data))
            return json.dumps(obj, ensure_ascii=False, indent=2)
        except Exception:
            text = _decode_text(data)
            return text or f"[unreadable json document: {filename}]"

    if ext == ".csv":
        text = _decode_text(data)
        sio = io.StringIO(text)
        reader = csv.reader(sio)
        lines: list[str] = []
        for row in reader:
            lines.append(" | ".join(str(c) for c in row))
        return "\n".join(lines).strip() or f"[empty csv document: {filename}]"

    if ext == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore[import-untyped]

            reader = PdfReader(io.BytesIO(data))
            pages = [(p.extract_text() or "").strip() for p in reader.pages]
            text = "\n\n".join(t for t in pages if t)
            return text or f"[pdf with no extractable text: {filename}]"
        except Exception as e:
            raise ValueError(f"Could not parse PDF '{filename}': {e}") from e

    if ext == ".docx":
        try:
            from docx import Document  # type: ignore[import-untyped]

            doc = Document(io.BytesIO(data))
            lines = [(p.text or "").strip() for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    lines.append(" | ".join((cell.text or "").strip() for cell in row.cells))
            text = "\n".join(t for t in lines if t)
            return text or f"[docx with no extractable text: {filename}]"
        except Exception as e:
            raise ValueError(f"Could not parse DOCX '{filename}': {e}") from e

    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook  # type: ignore[import-untyped]

            wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
            lines: list[str] = []
            for ws in wb.worksheets:
                lines.append(f"[sheet: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    if not row:
                        continue
                    vals = ["" if v is None else str(v) for v in row]
                    if any(v.strip() for v in vals):
                        lines.append(" | ".join(vals))
            text = "\n".join(lines).strip()
            return text or f"[xlsx with no extractable text: {filename}]"
        except Exception as e:
            raise ValueError(f"Could not parse XLSX '{filename}': {e}") from e

    if ext == ".xls":
        try:
            import xlrd  # type: ignore[import-untyped]

            book = xlrd.open_workbook(file_contents=data)
            lines: list[str] = []
            for sheet in book.sheets():
                lines.append(f"[sheet: {sheet.name}]")
                for rx in range(sheet.nrows):
                    row = [str(sheet.cell_value(rx, cx) or "") for cx in range(sheet.ncols)]
                    if any(c.strip() for c in row):
                        lines.append(" | ".join(row))
            text = "\n".join(lines).strip()
            return text or f"[xls with no extractable text: {filename}]"
        except Exception as e:
            raise ValueError(f"Could not parse XLS '{filename}': {e}") from e

    if ext == ".doc":
        raise ValueError("Legacy .doc is not supported in Azure local extractor. Convert to .docx.")

    raise ValueError(f"Unsupported file type for Azure RAG extraction: {ext or filename}")


class AzureRagService:
    async def get_or_create_search_index_id(
        self,
        agent_id: str,
        *,
        field_id: str | None,
        source: str,
        agent_settings: dict[str, Any] | None = None,
    ) -> str:
        scope_field_id = normalize_knowledge_field_id(field_id, source=source)
        existing = await knowledge_repository.get_backend_store_id(agent_id, field_id=scope_field_id)
        if existing:
            await ensure_agent_index(existing, agent_settings=agent_settings)
            return existing
        index_name = _safe_index_name(agent_id, scope_field_id)
        await ensure_agent_index(index_name, agent_settings=agent_settings)
        await knowledge_repository.upsert_vector_store(agent_id, index_name, field_id=scope_field_id)
        return index_name

    async def ingest_bytes(
        self,
        *,
        agent_id: str,
        data: bytes,
        filename: str,
        mime_type: str | None,
        source: str,
        field_id: str | None = None,
        thread_id: str | None = None,
        agent_settings: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        scope_field_id = normalize_knowledge_field_id(field_id, source=source)
        index_name = await self.get_or_create_search_index_id(
            agent_id,
            field_id=scope_field_id,
            source=source,
            agent_settings=agent_settings,
        )
        doc_id = str(uuid.uuid4())
        blob_uri = upload_archive_bytes(
            data,
            agent_id=agent_id,
            field_id=scope_field_id,
            artifact_ref=doc_id,
            filename=filename,
            content_type=mime_type,
        )
        content = _bytes_to_text(data, mime_type, filename).strip()
        if not content:
            raise ValueError(f"No extractable content from '{filename}'")
        content = _truncate_content_for_azure_index(content, filename=filename)
        await upload_documents(
            index_name=index_name,
            docs=[
                {
                    "id": doc_id,
                    "agent_id": agent_id,
                    "filename": filename,
                    "content": content,
                    "source": source,
                    "field_id": scope_field_id,
                    "thread_id": thread_id or "",
                    "blob_url": blob_uri or "",
                }
            ],
            agent_settings=agent_settings,
            # get_or_create_search_index_id already called ensure_agent_index
            ensure_index_exists=False,
        )
        logger.info(
            "Azure RAG: uploaded document id=%s index=%s filename=%r chars=%s",
            doc_id,
            index_name,
            filename,
            len(content),
        )
        if app_settings.azure_ai_search_verify_read_after_write:
            try:
                await get_document(
                    index_name,
                    doc_id,
                    agent_settings=agent_settings,
                )
                logger.info(
                    "Azure RAG: read-after-write GET succeeded (document is searchable), id=%s",
                    doc_id,
                )
            except Exception as e:
                logger.warning(
                    "Azure RAG: read-after-write GET failed id=%s index=%s — portal may disagree: %s",
                    doc_id,
                    index_name,
                    e,
                )
        await knowledge_repository.insert_knowledge_asset(
            agent_id=agent_id,
            source=source,
            field_id=scope_field_id,
            thread_id=thread_id,
            filename=filename,
            gcs_uri=blob_uri,
            provider_file_id=doc_id,
            backend_store_id=index_name,
        )
        return {
            "backend": "azure_ai_search",
            "backend_store_id": index_name,
            "vector_store_id": index_name,
            "provider_file_id": doc_id,
            "openai_file_id": doc_id,
            "gcs_uri": blob_uri,
            "filename": filename,
        }

    async def search_knowledge(
        self,
        *,
        agent_id: str,
        query: str,
        max_results: int = 8,
        field_id: str | None = None,
        agent_settings: dict[str, Any] | None = None,
    ) -> str:
        scope_field_id = normalize_knowledge_field_id(field_id, source="settings")
        index_name = await self.fetch_backend_store_id(agent_id, field_id=scope_field_id)
        if not index_name:
            return ""
        rows = await search_documents(
            index_name=index_name,
            query=query,
            max_results=max_results,
            agent_id=agent_id,
            field_id=scope_field_id,
            agent_settings=agent_settings,
        )
        parts: list[str] = []
        for r in rows:
            t = (r.get("content") or "").strip()
            if t:
                parts.append(t)
        return "\n\n---\n\n".join(parts[:max_results]) if parts else ""

    async def fetch_backend_store_id(self, agent_id: str, *, field_id: str | None = None) -> str | None:
        scope_field_id = normalize_knowledge_field_id(field_id, source="settings")
        return await knowledge_repository.get_backend_store_id(agent_id, field_id=scope_field_id)

    async def delete_ingested_file(
        self,
        *,
        agent_id: str,
        provider_file_id: str,
        backend_store_id: str | None = None,
        agent_settings: dict[str, Any] | None = None,
    ) -> None:
        row = await knowledge_repository.get_by_provider_file_id(agent_id, provider_file_id)
        index_name = str((row or {}).get("vector_store_id") or backend_store_id or "")
        blob_url = (row or {}).get("gcs_uri")

        if index_name:
            try:
                logger.info(
                    "Azure RAG: deleting index=%s doc_id=%r agent=%s row_in_db=%s",
                    index_name,
                    provider_file_id,
                    agent_id,
                    "yes" if row else "no",
                )
                key_present_before: bool | None = None
                try:
                    await get_document(
                        index_name,
                        provider_file_id,
                        agent_settings=agent_settings,
                    )
                    key_present_before = True
                except ResourceNotFoundError:
                    key_present_before = False
                    logger.warning(
                        "Azure RAG: DELETE is a no-op for key %r — that id is not in the index. "
                        "If the portal still shows a row, open it and compare the document key to this value.",
                        provider_file_id,
                    )
                await delete_documents(
                    index_name=index_name,
                    doc_ids=[provider_file_id],
                    agent_settings=agent_settings,
                )
                still_there = False
                try:
                    await get_document(
                        index_name,
                        provider_file_id,
                        agent_settings=agent_settings,
                    )
                    still_there = True
                except ResourceNotFoundError:
                    pass
                if key_present_before is True and still_there:
                    logger.error(
                        "Azure RAG: GET still returns the document after delete batch (key=%r index=%s)",
                        provider_file_id,
                        index_name,
                    )
                elif key_present_before is True and not still_there:
                    logger.info(
                        "Azure RAG: delete verified by GET (removed) index=%s doc_id=%r",
                        index_name,
                        provider_file_id,
                    )
                if key_present_before is False or still_there:
                    remaining = await search_documents(
                        index_name=index_name,
                        query="*",
                        max_results=50,
                        agent_id=agent_id,
                        field_id=None,
                        agent_settings=agent_settings,
                    )
                    if remaining:
                        logger.warning(
                            "Azure RAG: index %r still has %d document(s) for this agent; ids=%s filenames=%s",
                            index_name,
                            len(remaining),
                            [x.get("id") for x in remaining],
                            [x.get("filename") for x in remaining],
                        )
            except Exception as e:
                logger.warning(
                    "Azure Search delete failed (agent=%s index=%s doc=%s): %s",
                    agent_id,
                    index_name,
                    provider_file_id,
                    e,
                )

        delete_archive_uri(blob_url if isinstance(blob_url, str) else None)
        await knowledge_repository.delete_by_provider_file_id(agent_id, provider_file_id)


azure_rag_service = AzureRagService()
